import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Header Title
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run("IEEE COMPUTER SOCIETY STUDENT BRANCH CHAPTER\nMadhav Institute of Technology and Science (MITS), Gwalior")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.space_before = Pt(6)
    p_banner.paragraph_format.space_after = Pt(2)
    r_banner = p_banner.add_run("HACKMATRIX 2026 — ROUND 2")
    r_banner.font.bold = True
    r_banner.font.size = Pt(22)
    r_banner.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("CHOOSE YOUR REALITY. BUILD YOUR FUTURE.\nOFFICIAL PROJECT DOCUMENTATION")
    r_sub.font.bold = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    # Section Helper
    def add_sec_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title_text)
        r.font.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    # 1. Metadata Table
    add_sec_heading("PROJECT & TEAM METADATA")
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        ("Team Name", "COLDSTACK"),
        ("Team Leader Name & Contact", "Pranesh M S (praneshms2007@gmail.com | +91 9488344710)"),
        ("Team Members", "Pranesh M S, Avinash A S, Kannan S, Yashwanth"),
        ("Problem Statement", "PS-B2B-01 (B2B Industrial Market Fragmentation & Supply Chain Inefficiency)"),
        ("Event Name", "HackMatrix 2026 - Round 2"),
        ("Project Title", "IndusMate — Unified Sealed B2B Industrial Bidding Engine & AI Symbiosis"),
    ]

    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.8)
        set_cell_background(cell_k, "F8FAFC")
        set_cell_margins(cell_k, 80, 80, 120, 120)
        set_cell_margins(cell_v, 80, 80, 120, 120)

        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)

        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)

    # 2. Links Table
    add_sec_heading("OFFICIAL PROJECT LINKS")
    links_table = doc.add_table(rows=4, cols=2)
    links_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    links_table.autofit = False

    links_data = [
        ("GitHub Repository Link (Public)", "https://github.com/praneshMS-2007/Indusmate"),
        ("Live Deployed Link (Vercel)", "https://indusmate.vercel.app"),
        ("Demo Video Drive Link", "https://drive.google.com/drive/folders/1uNk-gQDxbypHJ-23VQl4YRwfi5fk1_gH"),
        ("Presentation Deck Canva Link", "https://canva.link/rbn90jqm4jr5e5m"),
    ]

    for i, (k, v) in enumerate(links_data):
        row = links_table.rows[i]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.8)
        set_cell_background(cell_k, "F8FAFC")
        set_cell_margins(cell_k, 80, 80, 120, 120)
        set_cell_margins(cell_v, 80, 80, 120, 120)

        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)

        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        rv.font.underline = True
        rv.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    # 3. Summary
    add_sec_heading("PROJECT SUMMARY")
    p_sum1 = doc.add_paragraph()
    p_sum1.paragraph_format.space_after = Pt(6)
    p_sum1.add_run("IndusMate is a production-grade, unified B2B industrial trading and byproduct symbiosis platform designed specifically for Indian manufacturing MSMEs, fleet transporters, raw material suppliers, and recyclers.")

    p_sum2 = doc.add_paragraph()
    p_sum2.paragraph_format.space_after = Pt(6)
    p_sum2.add_run("Built on the core insight that a raw material lot, an idle machine-hour, a technician's shift, a waste byproduct stream, and a truck's return leg are all instances of the same abstract object — a listable capacity with a specification, a location, a time window, and an undecided price — IndusMate unifies five distinct industrial markets onto one single database model and one universal deal state machine. It introduces strict server-side sealed anonymous bidding to prevent price cartelization and middleman markups (12–18%), alongside an AI-powered byproduct symbiosis engine using Google Gemini 2.5 to match industrial waste streams directly with recycling feedstocks.")

    # 4. Problem Being Solved
    add_sec_heading("PROBLEM BEING SOLVED")
    problems = [
        ("Broken Price Discovery & Middleman Markups: ", "Freight rates and raw material prices are set arbitrarily by intermediaries, imposing 12–18% unnecessary markups and inflating MSME operational costs."),
        ("Public Bid Leakage & Price Cartelization: ", "Traditional open tenders leak tenderer amounts before closing. Suppliers refrain from offering their true lowest prices because public disclosure exposes their commercial margins to rival buyers."),
        ("Unmapped Industrial Waste & Byproducts: ", "68% of industrial waste (fly ash, slag, chemical sludge, spent catalyst) is landfilled because factories search materials by brand name rather than chemical composition, turning sellable resources into disposal costs."),
        ("Unutilized Logistics Fleet Backhaul: ", "Transporters return with empty trucks on 35% of interstate routes due to lack of real-time return-leg freight visibility.")
    ]
    for b_title, b_desc in problems:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        rt = bp.add_run(b_title)
        rt.font.bold = True
        bp.add_run(b_desc)

    # 5. USP
    add_sec_heading("UNIQUE SELLING POINT (USP)")
    usps = [
        ("One Engine, Five Markets: ", "Exactly one Listing database model and one state machine power all 5 markets (Freight, Raw Materials, Byproducts, Equipment, Labour). Adding a 6th market requires zero structural code changes — only a spec configuration payload."),
        ("Strict Server-Side Sealed Anonymity (maskBid()): ", "Bidders never see rival amounts. Listing owners view pseudonymous handles and reputation metrics (Transporter #4471 · 4.7/5 · 128 deals). Private identities (GSTIN, legal name, phone) are released strictly upon deal ACCEPTED."),
        ("AI Waste-to-Resource Symbiosis Engine: ", "Uses Google Gemini 2.5 LLM to analyze chemical & physical specifications of industrial byproducts and match selling factories directly with buying recyclers."),
        ("Enterprise Admin KYC & Real-Time Audit Logs: ", "Full multi-document KYC verification flow with dedicated administrative queue (/admin/kyc) and transparent audit logging (/admin/logs)."),
        ("Sub-100ms Request-Level Performance: ", "Powered by Next.js 16 App Router, React 19, React.cache() query memoization, WebP compressed assets, Vercel Mumbai (bom1) edge functions, and Supabase PgBouncer connection pooling.")
    ]
    for b_title, b_desc in usps:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        rt = bp.add_run(b_title)
        rt.font.bold = True
        bp.add_run(b_desc)

    # 6. Key Features
    add_sec_heading("KEY FEATURES")
    features = [
        "Sealed-Bid Reverse Auctions (price competes down) & Forward Auctions (price competes up).",
        "Role-Aware Bento Dashboards for Manufacturers, Transporters, Suppliers, and Recyclers.",
        "AI Industrial Byproduct Symbiosis Matcher for automated waste stream valuation & counterparty matching.",
        "Platform Admin KYC Verification Queue (/admin/kyc) & Historical Audit Logs (/admin/logs).",
        "Interactive Leaflet OpenStreetMap cluster map for regional logistics and corridor tracking.",
        "Zero-latency useFormStatus submit buttons preventing duplicate form execution."
    ]
    for feat in features:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        bp.add_run(feat)

    # 7. Tech Stack Table
    add_sec_heading("TECHNICAL STACK")
    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_table.autofit = False

    t_headers = ["Layer", "Technologies Used", "Implementation Role"]
    hdr_row = tech_table.rows[0]
    for j, h_text in enumerate(t_headers):
        cell = hdr_row.cells[j]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(10)

    tech_rows = [
        ("Frontend", "Next.js 16 (App Router), React 19, TypeScript, Vanilla CSS + Tailwind v4", "Mobile-first responsive UI, Server Components for identity masking"),
        ("Backend", "Next.js Server Actions, Web API Route Handlers, Node.js", "Type-safe server actions, RESTful document APIs"),
        ("Database", "PostgreSQL (Supabase) + Prisma 6 ORM", "Request-level React.cache() memoization, PgBouncer pooling"),
        ("AI / LLM", "Google Gemini 2.5 Flash API (@google/genai)", "Automated byproduct waste stream matching & economic valuation"),
        ("Auth & Security", "Auth.js (NextAuth v5), Bcrypt, HTTP-only JWT Cookies", "Credentials auth, RBAC guards, Admin KYC verification"),
        ("Deployment", "Vercel Cloud Serverless (bom1 region) + Supabase Managed DB", "Sub-100ms global production edge hosting"),
    ]

    col_widths = [Inches(1.2), Inches(2.8), Inches(3.0)]

    for idx, (layer, tech, role) in enumerate(tech_rows, 1):
        row = tech_table.rows[idx]
        vals = [layer, tech, role]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.width = col_widths[j]
            set_cell_margins(cell, 80, 80, 120, 120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9.5)

    # 8. Future Scope
    add_sec_heading("FUTURE SCOPE & ROADMAP")
    roadmap = [
        ("Q3 2026 — Real-Time Counter-Offers & GSTN API: ", "Integrate WebSockets for live counter-bidding and automated government GSTN verification."),
        ("Q4 2026 — Payment Escrow Integration: ", "Integrate Razorpay / Cashfree escrow gateways to hold funds until milestone settlement."),
        ("Q1 2027 — Logistics Ecosystem Integration: ", "Connect directly with FASTag and NIC e-Way Bill portals for automated freight tracking."),
        ("2027+ — Industrial Hub Pilot & Scaling: ", "Deployment across Pithampur, Malanpur, and Mandideep industrial corridors, expanding to SAARC cross-border trade.")
    ]
    for r_title, r_desc in roadmap:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        rt = bp.add_run(r_title)
        rt.font.bold = True
        bp.add_run(r_desc)

    # 9. Team & Execution
    add_sec_heading("TEAM & EXECUTION STATEMENT")
    p_team = doc.add_paragraph()
    p_team.paragraph_format.space_after = Pt(4)
    r_t1 = p_team.add_run("Team COLDSTACK: ")
    r_t1.font.bold = True
    p_team.add_run("Pranesh M S (Lead / Full-Stack), Avinash A S (Frontend & Design Systems), Kannan S (Backend & Database Architecture), Yashwanth (QA & Testing).")

    p_stmt = doc.add_paragraph()
    p_stmt.paragraph_format.space_after = Pt(6)
    r_stmt = p_stmt.add_run("\"Combining deep expertise in full-stack Next.js web engineering, scalable database architecture, AI integration, and domain knowledge of Indian industrial supply chains to deliver a secure, frictionless B2B marketplace.\"")
    r_stmt.font.italic = True

    docx_path = r"C:\Users\Pranesh\Downloads\HACKMATRIX_PROJECT_DOCUMENTATION_INDUSMATE.docx"
    doc.save(docx_path)
    print("DOCX build successful:", docx_path)

if __name__ == "__main__":
    create_document()
